#!/usr/bin/env python3
"""Create polished Word field guides for the enterprise portfolio sections."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MA_OUT = ROOT / "enterprise-programs" / "01-ma-it-integration" / "M_AND_A_IT_INTEGRATION_PLAYBOOKS.docx"
HW_OUT = ROOT / "enterprise-programs" / "02-hardware-refresh" / "HARDWARE_REFRESH_FOUR_PLAYBOOKS.docx"

NAVY = "082E23"
TEAL = "0F9F8B"
BLUE = "145B7A"
INK = "142132"
MUTED = "5D6877"
PALE = "E8F6F1"
PALE_BLUE = "E9F3F7"
AMBER = "FFF0C7"
WHITE = "FFFFFF"
LINE = "D9E2E5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_rows(table) -> None:
    repeat_table_header(table.rows[0])
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (("Title", 30, WHITE), ("Heading 1", 19, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 10.5, TEAL)):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        # Give display headings enough breathing room when the next object is a
        # table caption. LibreOffice otherwise lets the caption ride into the
        # heading's descender box on some pages.
        style.paragraph_format.space_after = Pt(10 if name == "Heading 1" else 5)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.text = short_title.upper()
    p.style = styles["Caption"]
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(7.8)
    p.paragraph_format.space_after = Pt(0)
    footer = section.footer
    f = footer.paragraphs[0]
    f.add_run("Anthony Durham · Experience-informed fictional case study · ")
    f.runs[0].font.size = Pt(8)
    f.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(f)


def add_cover(doc: Document, label: str, title: str, subtitle: str, methods: str, stats: list[tuple[str, str]]) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=720, start=720, bottom=640, end=720)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label.upper())
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("75E1D1")
    p.paragraph_format.space_after = Pt(16)

    p = cell.add_paragraph()
    p.style = doc.styles["Title"]
    p.add_run(title)
    p.paragraph_format.space_after = Pt(12)

    p = cell.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor.from_string("DDF4F0")
    p.paragraph_format.space_after = Pt(18)

    p = cell.add_paragraph(methods)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string(WHITE)
    p.paragraph_format.space_after = Pt(22)

    stat_table = cell.add_table(rows=1, cols=len(stats))
    stat_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, (value, caption) in enumerate(stats):
        stat_cell = stat_table.cell(0, index)
        set_cell_shading(stat_cell, "0D4938")
        set_cell_margins(stat_cell, top=120, start=120, bottom=120, end=120)
        sp = stat_cell.paragraphs[0]
        sr = sp.add_run(value)
        sr.font.size = Pt(15)
        sr.font.bold = True
        sr.font.color.rgb = RGBColor.from_string("83E5D6")
        sp = stat_cell.add_paragraph(caption)
        sp.runs[0].font.size = Pt(7.8)
        sp.runs[0].font.color.rgb = RGBColor.from_string("D6EAE5")

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("PORTFOLIO GUIDE · AUGUST 2026")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("75E1D1")

    doc.add_page_break()
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def add_callout(doc: Document, heading: str, text: str, fill: str = PALE) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_margins(c, top=130, start=150, bottom=130, end=150)
    p = c.paragraphs[0]
    r = p.add_run(heading)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = c.add_paragraph(text)
    p.runs[0].font.size = Pt(8.8)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int = 1, page_break: bool = False) -> None:
    p = doc.add_heading(text, level=level)
    if page_break:
        # A page-break-before setting keeps the heading and following caption
        # in the same layout flow; a separate break paragraph can paint over
        # the first line in LibreOffice's PDF renderer.
        p.paragraph_format.page_break_before = True
    keep_with_next(p)


def add_paragraph(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(item, style=style)
        p.paragraph_format.space_after = Pt(2.5)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.size = Pt(9)
    keep_with_next(p)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        pr = cell.paragraphs[0]
        run = pr.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(8)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            if row_index % 2:
                set_cell_shading(cell, "F7FAFB")
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(7.8)
            run.font.color.rgb = RGBColor.from_string(INK)
            if widths:
                cell.width = Inches(widths[index])
    set_repeat_table_rows(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_reuse_box(doc: Document, fields: list[str]) -> None:
    add_callout(doc, "Reusable template fields", " · ".join(fields), PALE_BLUE)


def build_ma_guide() -> None:
    doc = Document()
    configure_document(doc, "M&A IT Integration Playbooks")
    add_cover(
        doc,
        "Enterprise transformation portfolio",
        "M&A IT Integration Playbooks",
        "A complete, anonymized program from clean-team discovery and Day 1 through secure convergence, decommissioning, and benefits ownership.",
        "Agile · Predictive / Waterfall — PMBOK Guide Eighth Edition aligned",
        [("680", "workers"), ("745", "endpoints"), ("142", "applications"), ("$1.86M", "modeled authorization")],
    )

    add_heading(doc, "How to read this portfolio", 1)
    add_callout(doc, "Truth boundary", "Company ABC and Company XYZ are fictional. All project dates, quantities, costs, performance, and decisions are scenario assumptions. Verified experience is separated from case results; product and standards references are current primary-source guidance as of August 25, 2026.", AMBER)
    add_table(doc, "Table 1. Evidence classes", ["Class", "Meaning", "Example"], [
        ["Verified experience", "Supported by supplied resume/profile", "Nine M&A endpoint integrations"],
        ["Public source", "Authoritative current guidance", "PMI PMBOK 8; NIST CSF 2.0"],
        ["Scenario assumption", "Fictional planning input", "680 workers; $1.86M"],
        ["Derived calculation", "Arithmetic from disclosed inputs", "665 / 680 = 97.8%"],
        ["Unknown", "Requires discovery, approval, test, contract, or source record", "Mailbox volume; legal-hold population"],
    ], [1.45, 2.55, 2.65])
    add_paragraph(doc, "The website contains the full artifact library and filterable workbook. This guide is the executive field manual: it shows the control architecture, method differences, key registers, and decision rules.")
    add_reuse_box(doc, ["deal thesis", "close date", "clean-team rules", "headcount/sites/devices/apps", "target state", "success measures", "method rationale"])

    add_heading(doc, "Scenario and outcome architecture", 1, page_break=True)
    add_table(doc, "Table 2. Scenario at a glance", ["Dimension", "Filled example", "Acceptance truth"], [
        ["Day 1", "680 workers across three offices and remote", "665 fully ready; 15 owned exceptions"],
        ["Identity/endpoints", "Company ABC target; temporary bridge; 745 endpoints", "Unique identity, MFA, device path, least privilege, logs"],
        ["Network", "Three sites, two clouds, four subnet collisions", "No unresolved overlap; restricted, redundant, observable flows"],
        ["Collaboration", "Google Workspace, Slack, Asana, Jira", "Counts, samples, permissions, holds, integrations, business acceptance"],
        ["Applications", "142 discovered; 37 overlaps; 22 targeted", "Owner, criticality, disposition, dependency, contract and evidence"],
        ["Financial", "$1.860M baseline; $1.782M forecast", "Planning assumption; Finance validates actual/realized value"],
    ], [1.3, 2.45, 2.9])
    add_heading(doc, "Experience that informs the design", 2)
    add_bullets(doc, [
        "Nine enterprise M&A endpoint integrations covering standards, access, onboarding, tooling, support, and continuity.",
        "Digital-workplace leadership across 37,000 Macs and 300,000 Windows endpoints with modern management and telemetry.",
        "Collaboration-platform work spanning Google Workspace, Slack, Box, Microsoft 365, Teams, OneDrive, Jira, and Confluence.",
        "A 45-person portfolio, $14 million financial responsibility, vendor savings, security/audit partnership, endpoint standards, and cross-functional delivery.",
    ])
    add_callout(doc, "What this does not claim", "The modeled Company ABC/XYZ project, quantities, dates, costs, migration results, and closeout are not presented as a prior employer engagement.")

    add_heading(doc, "Agile playbook", 1, page_break=True)
    add_paragraph(doc, "Product Goal: acquired workers perform priority work securely on Day 1, then Company ABC progressively converges technology and removes temporary risk and cost with measured acceptance.")
    add_table(doc, "Table 3. Agile operating model", ["Control", "Filled example", "Evidence"], [
        ["Cadence", "Ten two-week Sprints; reviews, retrospectives, twice-weekly dependencies", "Sprint goal, accepted Increment, decision log"],
        ["Releases", "Foundation; Day 1; migration waves; convergence/closure", "Release outcome and shared gate"],
        ["Backlog", "Outcome, owner, priority, acceptance, risk, dependency, evidence", "One ordered integration backlog"],
        ["Quality", "Definition of Ready/Done includes security, data, business, operations", "No tool-success-only completion"],
        ["Metrics", "Outcome burn-up, cycle time, throughput, defects, blockers, budget", "Three Sprints: 176/186 forecast points; P85 6.9 days"],
    ], [1.25, 2.7, 2.7])
    add_heading(doc, "Agile gate tension", 2)
    add_paragraph(doc, "Legal close, Day 1, security/privacy approval, contract dates, and destructive decommission actions are fixed constraints. Discovery and solution work are adaptive. The backlog may reorder work; it may not delete control evidence to protect velocity.")
    add_reuse_box(doc, ["Product Goal", "value streams", "Sprint length/events", "release map", "backlog schema", "DoR/DoD", "metrics", "fixed constraints"])

    add_heading(doc, "Predictive / Waterfall playbook", 1, page_break=True)
    add_paragraph(doc, "PMBOK Guide Eighth Edition is delivery-approach neutral. This playbook is called Predictive / Waterfall because it uses formal scope, WBS, schedule and cost baselines, sequenced stage gates, integrated change control, quality assurance, and earned value reporting.")
    add_table(doc, "Table 4. PMBOK 8 application", ["Element", "Portfolio application", "Primary evidence"], [
        ["Six principles", "Holistic view, value, embedded quality, accountability, sustainability, empowered teams", "Charter, acceptance, governance, closure"],
        ["Seven domains", "Governance, scope, schedule, finance, stakeholders, resources, risk", "Baselines, owners, measures and registers"],
        ["Process groups", "Initiating, Planning, Executing, Monitoring/Controlling, Closing", "Lifecycle and gate sequence"],
        ["Tailoring", "Formal legal/security/data/decommission gates; rolling-wave detail for unknowns", "Tailoring statement and change thresholds"],
    ], [1.3, 2.9, 2.45])
    add_table(doc, "Table 5. Predictive performance checkpoint", ["Measure", "Evidence", "Interpretation"], [
        ["SPI", "0.97 at Week 12", "Schedule watch; Week 24 unchanged"],
        ["CPI", "1.02", "Cost favorable; never substitutes for acceptance"],
        ["EAC / VAC", "$1.782M / $78K favorable", "Fictional; update from source financial records"],
        ["Critical control", "Jira production cohort held", "Quality gate prevents premature release"],
    ], [1.3, 2.2, 3.15])
    add_reuse_box(doc, ["tailoring", "WBS/control accounts", "IMS/critical path", "cost baseline/EVM", "stage gates", "quality", "change thresholds", "closure"])

    add_heading(doc, "Shared control pack", 1, page_break=True)
    add_table(doc, "Table 6. Shared artifacts and accountable questions", ["Artifact", "Question it answers", "Completion evidence"], [
        ["Charter/governance", "What is authorized; who decides?", "Signed outcomes, limits, forums, thresholds"],
        ["Day 1 plan", "Can every worker perform priority work?", "Roster/persona/device/access/support evidence"],
        ["Identity/access", "Is the identity bridge controlled and temporary?", "Unique IDs, least privilege, privileged review, expiry"],
        ["Network/infrastructure", "Can required traffic flow safely and observably?", "Collision, DNS, route, rule, failover and log tests"],
        ["Collaboration migration", "Is content usable, permitted, retained and integrated?", "Counts, samples, permissions, holds, bots/webhooks"],
        ["Application rationalization", "Does every app have a safe target disposition?", "142/142 owner, decision, dependency and exit evidence"],
        ["Security/privacy", "Are risk, data and legal controls embedded?", "NIST CSF outcomes, data approvals, risk acceptance"],
        ["Cutover/hypercare", "Can we stop, recover and hand off?", "Time-stamped runbook, triggers, rollback and service acceptance"],
    ], [1.45, 2.6, 2.6])

    add_heading(doc, "Day 1 readiness and support", 1, page_break=True)
    add_table(doc, "Table 7. Modeled Day 1 readiness", ["Cohort", "Population", "Fully ready", "Exceptions", "Control"], [
        ["HQ Site A", "280", "272", "8", "Loaners + temporary approved access"],
        ["Office Site B", "190", "186", "4", "Site runner + spares"],
        ["Office Site C", "110", "107", "3", "Concierge appointments"],
        ["Remote", "100", "100", "0", "Courier + virtual support"],
        ["Total", "680", "665 (97.8%)", "15", "Every exception has owner/workaround/expiry"],
    ], [1.35, 0.8, 1.15, 0.8, 2.55])
    add_bullets(doc, [
        "Freeze and reconcile HR roster at T-10, T-3 and T-0; resolve duplicates, contractors, managers, personas, sites, devices and special access.",
        "Provision destination identities disabled; validate MFA, privileged roles, service identities, licenses and recovery before release.",
        "Stage devices and instructions by site/remote logistics; test standard, executive/delegate, developer/admin, shared, contractor and accessibility personas.",
        "Open a staffed command center; measure login, compliance, collaboration, priority app, incident, workaround and user-impact evidence at T+2/T+8.",
    ])
    add_callout(doc, "Stop rule", "Hold a cohort if roster-to-identity mismatch exceeds 1%, a privileged path is unapproved, a Tier 0 workflow fails, or support cannot receive and route cases. Preserve the approved source path; rollback is not deletion.")

    add_heading(doc, "Migration workstreams and acceptance", 1, page_break=True)
    add_table(doc, "Table 8. Platform migration control", ["Platform", "Scope", "High-risk edge cases", "Acceptance"], [
        ["Google Workspace", "610 users; Gmail, Calendar, Drive, Shared Drives, Groups", "Delegates, sharing, ownership, holds, links, unsupported items", "99.5%+ counts plus samples, permissions and business tests"],
        ["Slack", "520 users; 286 channels", "Private/shared channels, guests, history, files, apps, bots, webhooks", "Membership/history/integrations/retention accepted"],
        ["Asana", "280 users", "Owners, dates, custom fields, archives and privacy", "Project/task/field reconciliation"],
        ["Jira", "190 users", "Email identity, groups, permissions, apps, filters, boards, automation", "User/group-first mapping and workflow regression"],
        ["Applications", "142 discovered", "Shadow IT, contracts, data, APIs, ownerless apps", "Owner, disposition, migration/archive, contract and decommission evidence"],
    ], [1.15, 1.85, 2.25, 2.15])
    add_paragraph(doc, "A migration batch is not accepted on a tool-success message. The accountable owner reviews object counts, defined exclusions, representative samples, permissions, retention/legal hold, critical integrations, searchability and business workflows. The source remains available/read-only for the approved retention window.")

    add_heading(doc, "Governance, risk, and financial control", 1, page_break=True)
    add_table(doc, "Table 9. Top RAID items", ["ID", "Exposure", "Score", "Response"], [
        ["MA-R-001", "Identity duplicates/stale contractors", "20", "T-10/T-3/T-0 reconciliation and exception path"],
        ["MA-R-002", "Subnet overlap disrupts routing/DNS", "15", "NAT/renumber, isolated test, route withdrawal"],
        ["MA-R-003", "Collaboration permissions/retention/integrations differ", "15", "Identity-first mapping, legal review, pilot and validation"],
        ["MA-I-004", "15 Day 1 worker exceptions", "15", "Workarounds, owners, expiry and daily burn-down"],
    ], [0.85, 2.55, 0.7, 2.9])
    add_table(doc, "Table 10. Budget control", ["Measure", "Modeled value", "Control"], [
        ["Authorized baseline", "$1.860M", "Scenario assumption"],
        ["Forecast at completion", "$1.782M", "Weekly ETC/EAC from work packages"],
        ["Forecast variance", "$78K favorable", "Protect reserve and quality gates"],
        ["Savings", "Not booked", "Finance validates contract/invoice stop; capacity is not cash"],
    ], [1.55, 1.55, 3.65])

    add_heading(doc, "Cutover, rollback, and hypercare", 1, page_break=True)
    add_table(doc, "Table 11. Gate and command path", ["Checkpoint", "Required evidence", "Decision"], [
        ["T-10", "Roster, inventory, design, vendor, communications, test status", "Continue / recover / de-scope"],
        ["T-3", "No Sev 1; Tier 0 tests; rollback; staffing", "Go / conditional go / hold"],
        ["T-0", "Backups/exports, freeze, monitoring, bridge, approvals", "Start cutover"],
        ["Batch validation", "Counts, samples, permissions, workflow, support", "Accept / retry / rollback cohort"],
        ["Hypercare exit", "Stable SLA, defects, knowledge, on-call, owner acceptance", "Transfer to operations"],
    ], [1.25, 3.65, 1.85])
    add_bullets(doc, [
        "Rollback is assessed per identity cohort, route/rule group, application, site and data batch—not as one all-or-nothing command.",
        "Triggers include security control failure, permission variance, reconciliation below 99.5%, identity failure above 2%, Tier 0 workflow failure, or inability to restore within the approved objective.",
        "Exit hypercare after five stable business days, no Sev 1, accepted service ownership, controlled aged defects and documented residual risk.",
    ])

    add_heading(doc, "Executive dashboard and decision rules", 1, page_break=True)
    add_table(doc, "Table 12. Illustrative checkpoint", ["Dimension", "Evidence", "Decision"], [
        ["Day 1", "665/680 fully ready; 674/680 authentication", "Close 15 exceptions by Day 3"],
        ["Migration", "Workspace pilot 99.7%; Slack/Jira integration defects open", "Hold affected production scope"],
        ["Applications", "142 discovered; five Tier 2 owners unresolved", "Assign/restrict before integration"],
        ["Schedule", "Agile Week 20 or Predictive Week 24 forecast", "Method-specific forecast; same acceptance"],
        ["Finance", "$1.782M forecast", "Favorable but low-confidence scenario value"],
    ], [1.25, 3.25, 2.25])
    add_callout(doc, "Management rule", "A dashboard is green only when its source register is current. Unknown denominator, unowned exception, missing approval, untested rollback or unvalidated destructive action is visible and blocks the affected decision.")

    add_heading(doc, "Artifact inventory and research", 1, page_break=True)
    add_table(doc, "Table 13. Repository package", ["Package", "Examples", "Format"], [
        ["Shared controls", "Charter; governance; Day 1; IAM; network; collaboration; apps; security; change; cutover; vendor", "Rendered web pages + Markdown"],
        ["Registers", "RAID; dependencies; RACI; app portfolio; readiness; validation; budget; sources", "Rendered tables + CSV"],
        ["Agile", "Playbook; roadmap; backlog; Sprint plan; DoR/DoD; metrics; status", "Web/Markdown/CSV"],
        ["Predictive", "Playbook; PMBOK 8 map; WBS; IMS; gates/change; quality; EVM; status", "Web/Markdown/CSV"],
        ["Control workbook", "Two dashboards and reusable source registers/formulas", "XLSX"],
    ], [1.35, 3.7, 1.7])
    add_paragraph(doc, "Primary-source research includes PMI PMBOK Guide Eighth Edition, PMI Process Groups and Agile Practice Guides, the 2026 PMP exam update, NIST CSF 2.0, Microsoft Entra multi-tenant and Azure VPN guidance, Google Workspace migration resources, Slack migration documentation, Atlassian migration guidance, and Asana import/export documentation.")
    add_reuse_box(doc, ["source title", "publisher", "URL", "access date", "portfolio use", "authority", "validation owner"])

    MA_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(MA_OUT)


def build_hardware_guide() -> None:
    doc = Document()
    configure_document(doc, "Hardware Refresh — Four Playbooks")
    add_cover(
        doc,
        "Digital workplace transformation portfolio",
        "Hardware Refresh — Four Playbooks",
        "A 360-device, four-site endpoint program from standards and procurement through user acceptance, secure retirement, service transition, and benefits.",
        "Kanban · Scrum · Predictive / Waterfall — PMBOK 8 aligned · Hybrid",
        [("360", "devices"), ("4", "sites"), ("18", "weeks"), ("$792K", "modeled authorization")],
    )

    add_heading(doc, "How to read this portfolio", 1)
    add_callout(doc, "Truth boundary", "Company ABC is fictional. All case dates, quantities, costs and results are scenario assumptions. Verified endpoint, digital-workplace and delivery-method experience is separated from modeled results. Product and standards references are current primary-source guidance as of August 25, 2026.", AMBER)
    add_table(doc, "Table 1. Experience-informed design", ["Verified experience", "Control informed", "Not claimed"], [
        ["37K Macs / 300K Windows modern management", "Inventory, build, compliance, telemetry, support", "The 360-device result is fictional"],
        ["Intune, SCCM/MECM, Jamf, ServiceNow, Tanium, BeyondTrust", "Provisioning, management, workflow, support, analytics", "Actual Company ABC licensing/architecture is unknown"],
        ["98% OS compliance; endpoint standards; audit partnership", "Release thresholds, policy, evidence, sanitization", "No certification or legal opinion"],
        ["Scrum/Kanban/program leadership", "Four tailored delivery systems", "Method metrics are modeled"],
    ], [2.05, 2.55, 2.1])
    add_reuse_box(doc, ["device/site/persona baseline", "method rationale", "standards", "budget/schedule", "success measures", "service/benefit owners"])

    add_heading(doc, "Program architecture and outcomes", 1, page_break=True)
    add_table(doc, "Table 2. Scenario scope", ["Dimension", "Filled example", "Acceptance"], [
        ["Devices", "270 laptops; 60 desktops; 30 phones", "360 final dispositions"],
        ["Sites", "Three offices plus remote/all-site cohort", "Manager roster, blackout and readiness sign-off"],
        ["Technology", "Intune/Entra target; SCCM/MECM approved existing-device path", "Versioned build, join/enrollment, security and support"],
        ["User/data", "76 apps; KFM/approved backup; specialized personas", "Task-based app/data/accessibility acceptance"],
        ["Old assets", "Return, custody, hold release, sanitization and disposition", "No unknown serial or evidence gap"],
        ["Financial", "$792K baseline; $748.4K forecast", "Planning assumption; reconcile PO/invoice/credits/reserve"],
    ], [1.2, 2.8, 2.8])
    add_table(doc, "Table 3. Illustrative checkpoint", ["Measure", "Result", "Interpretation"], [
        ["Deployed", "344 / 360", "16 controlled exceptions remain"],
        ["First-time-right", "336 / 344 = 97.7%", "Meets 97% target"],
        ["Data validation", "342 / 344 = 99.4%", "Two owned exceptions"],
        ["Compliant within 24h", "339 / 344 = 98.5%", "Five remediations"],
        ["Old devices returned", "334 / 344 = 97.1%", "Remote return action open"],
    ], [1.65, 1.8, 3.35])

    add_heading(doc, "Method selection", 1, page_break=True)
    add_table(doc, "Table 4. Four operating systems", ["Method", "Best fit", "Primary controls", "Watch-out"], [
        ["Kanban", "Variable ready demand and exceptions", "Workflow, WIP, SLEs, age, cycle time, throughput", "A board without pull/limits is only visualization"],
        ["Scrum", "Cross-functional team can deliver accepted wave Increments", "Product Goal, Backlog, Sprints, DoD, Review/Retro", "Do not count staged devices as Done"],
        ["Predictive", "Fixed quantities, site order, procurement and formal approvals", "WBS, IMS, baselines, gates, change, EVM", "Unknowns need rolling-wave detail, not false precision"],
        ["Hybrid", "Fixed compliance/supply guardrails plus variable user/app work", "Baseline/gates + pull/WIP + rolling waves", "One source of truth per control; no competing plans"],
    ], [1.0, 2.05, 2.35, 1.65])
    add_callout(doc, "Recommendation", "Hybrid is the modeled recommendation because supply, budget, security and site commitments are fixed, while application defects, user readiness, local data, appointments and remote returns vary. The other three playbooks remain complete for comparison.")

    add_heading(doc, "Kanban playbook", 1, page_break=True)
    add_table(doc, "Table 5. Definition of Workflow", ["State", "Exit policy", "WIP"], [
        ["Ready", "Roster/persona, hardware, app/data precheck and appointment feasible", "60"],
        ["Staging", "Build/security/asset test passes", "24"],
        ["Scheduled", "User/site communication and kit ready", "40"],
        ["Deploying", "Task/data/user acceptance", "12"],
        ["Hypercare", "T+5 measures pass or exception assigned", "30"],
        ["Return/Sanitize", "Disposition evidence accepted", "35"],
    ], [1.55, 4.35, 0.8])
    add_bullets(doc, [
        "Service classes: Expedite, Fixed Date, Standard and Intangible; one Expedite maximum.",
        "Review aging at 50% of SLE, swarm at 75%, escalate or reshape at 100%.",
        "Latest modeled flow: 42 devices/week, median 3.4 days, P85 6.2 days; stop starting and finish exception/return work.",
        "Done includes new-device acceptance and old-asset reconciliation; handoff alone is not the delivery point.",
    ])
    add_reuse_box(doc, ["workflow states", "entry/exit", "WIP", "classes", "SLEs", "cadences", "flow/outcome metrics", "policy experiments"])

    add_heading(doc, "Scrum playbook", 1, page_break=True)
    add_paragraph(doc, "Product Goal: by Sprint 9, all 360 assignments have an accepted secure replacement or approved final exception; old assets are reconciled; Service Operations owns the stable service.")
    add_table(doc, "Table 6. Scrum design", ["Element", "Filled example", "Evidence"], [
        ["Accountabilities", "Endpoint Service Owner; Agile Delivery Lead; cross-functional Developers", "Clear Product Owner/Scrum Master/Developers"],
        ["Sprints", "Nine two-week Sprints", "Sprint Goal, Sprint Backlog and usable Increment"],
        ["Increment", "Accepted cohort plus security/data/asset/support evidence", "Strict Definition of Done"],
        ["Progress", "344 accepted; 16 PBIs remain", "Incomplete work returns to Product Backlog"],
        ["Inspection", "Stakeholder Review; Retrospective improvement", "Task evidence and one measurable experiment"],
    ], [1.35, 3.0, 2.4])
    add_reuse_box(doc, ["Product Goal", "accountabilities", "Sprint events", "Product Backlog", "Sprint Goal", "Definition of Done", "release/wave gates", "outcome metrics"])

    add_heading(doc, "Predictive / Waterfall playbook", 1, page_break=True)
    add_paragraph(doc, "PMBOK Guide Eighth Edition remains approach-neutral. This version uses a predictive sequence because quantities, sites, budget, supply, standards and approvals can be baselined; exceptions receive rolling-wave detail inside controlled work packages.")
    add_table(doc, "Table 7. Predictive controls", ["Control", "Filled example", "Decision use"], [
        ["WBS", "Governance; supply; engineering; deployment; retirement; reserve", "Accepted work packages/control accounts"],
        ["IMS", "Week 1–18 with pilot and four waves", "Critical path and milestone forecast"],
        ["Gates", "Charter, baseline, build/pilot, scale, wave, retirement, close", "Formal authorization"],
        ["EVM", "SPI 0.97; CPI 1.03; EAC $748.4K", "Schedule/cost forecast with acceptance-based earning"],
        ["Change", "$15K/five-day PM/CCB threshold; steering above", "Protect baselines and controls"],
    ], [1.2, 3.0, 2.55])
    add_reuse_box(doc, ["PMBOK 8 mapping", "WBS", "IMS", "cost baseline/EVM", "stage gates", "quality", "change thresholds", "closure/benefits"])

    add_heading(doc, "Hybrid playbook", 1, page_break=True)
    add_table(doc, "Table 8. Integrated operating model", ["Layer", "Predictive guardrail", "Adaptive control"], [
        ["Program", "360 scope, $792K, Week 18, standards", "Reprioritize cohorts inside thresholds"],
        ["Supply/build", "PO, receipt, build release and pilot gates", "Daily pull; staging WIP 24; defect learning"],
        ["Deployment", "Four site milestones; T-5/T-1 go/no-go", "Two-week horizon; readiness/capacity pull"],
        ["Exceptions", "Risk/change authority; full disposition", "Dedicated lane, WIP 10, aging and swarming"],
        ["Reporting", "Baseline/forecast, quality, budget, RAID", "Flow/outcome metrics and retrospectives"],
    ], [1.2, 3.0, 2.55])
    add_paragraph(doc, "One source of truth per control: baseline for authorized scope/cost/date; wave register for released quantities and acceptance; board for current state; inventory for serial identity; RAID for exposure. Dashboards derive from those sources rather than narrative reconciliation.")
    add_reuse_box(doc, ["guardrails", "adaptive work", "planning horizons", "workflow/WIP", "wave gates", "decision matrix", "reconciliation", "metrics"])

    add_heading(doc, "Shared technical control pack", 1, page_break=True)
    add_table(doc, "Table 9. Technical and operational controls", ["Control", "Required evidence", "Stop condition"], [
        ["Device build", "Version, join/enrollment, apps, update, drivers, support", "Unknown/mismatched build or critical app"],
        ["Security", "Encryption, EDR, firewall, supported OS, compliance, least privilege", "Missing protection or unapproved exception"],
        ["Applications", "76 dispositions; Tier 0/1 package/license/task test/fallback", "Critical workflow failure"],
        ["User data", "Precheck, sync/backup, counts/samples, task/user acceptance", "Unknown local data, sync error or hold"],
        ["Asset", "New/old serial, user/site/wave, custody scans, authoritative state", "Unknown serial, loss or tamper"],
        ["Sanitization", "Approved method/program, validation, certificate, final disposition", "Hold not released or evidence/validation failure"],
        ["Support", "Queue, knowledge, staffing, spares, escalation, vendors, monitoring", "Cannot restore affected users"],
    ], [1.25, 3.75, 1.75])

    add_heading(doc, "Wave and user journey", 1, page_break=True)
    add_table(doc, "Table 10. Deployment waves", ["Wave", "Planned", "Deployed", "FTR", "Data", "24h compliance", "Returns", "Exceptions"], [
        ["Pilot", "20", "20", "19", "20", "20", "20", "0"],
        ["Site A", "95", "94", "92", "94", "93", "92", "1"],
        ["Site B", "90", "87", "85", "86", "86", "84", "3"],
        ["Site C", "75", "71", "69", "71", "70", "68", "4"],
        ["Remote/all", "80", "72", "71", "71", "70", "70", "8"],
        ["Total", "360", "344", "336", "342", "339", "334", "16"],
    ], [1.05, 0.72, 0.72, 0.64, 0.64, 1.0, 0.72, 0.85])
    add_bullets(doc, [
        "T-30 managers validate roster, blackout, persona and local needs; T-14 users receive appointment/preparation/data/app/accessibility instructions.",
        "T-5 prechecks classify Ready, controlled exception with new date, or not authorized; T-1 go/no-go verifies security, data, app, logistics, support and rollback.",
        "Day 0 verifies identity, old-asset custody, new-device issue, user tasks and acceptance; T+1/T+5 measures compliance, tickets, data, returns, satisfaction and lessons.",
    ])
    add_callout(doc, "Denominator rule", "A reschedule, leave, accessibility dependency, application defect or logistics delay remains in the 360-device denominator until it reaches an approved final disposition.")

    add_heading(doc, "Secure retirement and chain of custody", 1, page_break=True)
    add_paragraph(doc, "NIST SP 800-88 Rev. 2 is applied as enterprise sanitization-program guidance: classify data/media, select an organization-approved method/standard, authorize tools/vendors, preserve custody, validate effectiveness, retain evidence and manage reuse/disposal. The project does not invent a destructive technique.")
    add_table(doc, "Table 11. Custody states", ["State", "Evidence", "Gate"], [
        ["User handoff", "Old asset ID/serial, person, time, location, condition", "Identity and data release"],
        ["Secure storage", "Cage/container/seal and scan", "No unknown state"],
        ["Transport/processor", "Pickup, tracking, receipt and exception", "Vendor/chain acceptance"],
        ["Sanitization", "Method/tool/version, result, validator, exception", "Approved program and hold release"],
        ["Final disposition", "Certificate/reference and reuse/recycle/destruction", "Asset/records/security acceptance"],
    ], [1.25, 3.8, 1.7])
    add_callout(doc, "Quarantine rule", "Unknown identity or serial, missing hold release, suspected loss/tamper, failed sanitization validation or vendor evidence gap is quarantined and escalated. An invoice alone never closes custody.")

    add_heading(doc, "Governance, RAID, and financial control", 1, page_break=True)
    add_table(doc, "Table 12. Top RAID items", ["ID", "Exposure", "Score", "Response"], [
        ["HW-R-003", "Critical app/driver/VPN/peripheral failure", "16", "Persona pilot, versioned package, fallback, owner acceptance"],
        ["HW-R-002", "Local/application data not protected", "15", "Precheck, KFM/backup, validation, sanitize hold"],
        ["HW-I-005", "16 application/accessibility/leave/logistics exceptions", "15", "Named owner/date/workaround and specialized capacity"],
        ["HW-R-004", "Custody/sanitization evidence gap", "10", "Serial scans, secure storage, validation and quarantine"],
    ], [0.9, 2.65, 0.7, 2.8])
    add_table(doc, "Table 13. Budget", ["Measure", "Modeled value", "Rule"], [
        ["Baseline", "$792K", "Devices, accessories, shipping, staging, support, sanitization, change, reserve"],
        ["Forecast", "$748.4K", "Update from PO, receipt, invoice, labor, tax, credit and ETC"],
        ["Variance", "$43.6K favorable", "Do not trade quality/acceptance for cost"],
        ["Benefits", "Not yet realized", "Operational owner + Finance validate source evidence"],
    ], [1.35, 1.7, 3.95])

    add_heading(doc, "Service transition, closure, and benefits", 1, page_break=True)
    add_bullets(doc, [
        "Exit wave hypercare after five stable business days, no Sev 1, first-time-right at least 97%, compliance at least 98% within 24 hours, and open defects inside the agreed service objective.",
        "Service Operations accepts build, CMDB/inventory, knowledge, queue, monitoring, spares, vendor, escalation, on-call, KPIs and known risks.",
        "Close only when 360 of 360 new/old asset dispositions reconcile; data/holds and sanitization evidence close; POs/invoices/credits/warranties reconcile; risks/issues transfer; sponsor accepts outcome.",
        "Benefits continue after closure: supported-device rate, compliance, startup/app reliability, ticket/contact rate, satisfaction, repair/warranty and validated avoided cash cost.",
    ])
    add_reuse_box(doc, ["hypercare exit", "service acceptance", "asset reconciliation", "financial/vendor closure", "open-item transfer", "lessons/actions", "benefit owners/measures"])

    add_heading(doc, "Artifact inventory and research", 1, page_break=True)
    add_table(doc, "Table 14. Repository package", ["Package", "Method-specific artifacts", "Shared evidence"], [
        ["Kanban", "Workflow/policies, board, flow metrics, service review, status", "Full endpoint/data/asset control pack"],
        ["Scrum", "Product Goal/Backlog, Sprint plan, DoD, metrics, status", "Full endpoint/data/asset control pack"],
        ["Predictive", "PMBOK 8 map, WBS, IMS, gates/change, quality, EVM, status", "Full endpoint/data/asset control pack"],
        ["Hybrid", "Integrated governance, rolling-wave/flow, plan, metrics, status", "Full endpoint/data/asset control pack"],
        ["Workbook", "Four dashboards with formulas/charts", "RAID, inventory, waves, budget, custody, sources, templates"],
    ], [1.25, 3.25, 2.45])
    add_paragraph(doc, "Primary-source research includes PMI PMBOK Guide Eighth Edition, PMI Process Groups and Agile Practice Guides, the Scrum Guide, the Kanban Guide, NIST SP 800-88 Rev. 2, Microsoft Windows Autopilot, Intune compliance, OneDrive Known Folder Move and Endpoint Analytics guidance.")

    HW_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(HW_OUT)


def main() -> None:
    build_ma_guide()
    build_hardware_guide()
    print(f"Built {MA_OUT.relative_to(ROOT)}")
    print(f"Built {HW_OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

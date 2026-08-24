#!/usr/bin/env python3
"""Build four reusable robotics PM Word templates with fixed, auditable geometry."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "pm-operating-system" / "word"

# compact_reference_guide preset, resolved token map.
PAGE_WIDTH = 8.5
PAGE_HEIGHT = 11.0
MARGIN = 1.0
HEADER_DISTANCE = 0.492
FOOTER_DISTANCE = 0.492
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
MUTED = "5B6670"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run(run, *, size=11, bold=False, color="000000", italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Table Title" not in doc.styles:
        style = doc.styles.add_style("Table Title", WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles["Table Title"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.keep_with_next = True
    if "Evidence Note" not in doc.styles:
        style = doc.styles.add_style("Evidence Note", WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles["Evidence Note"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)


def new_document(title: str, subtitle: str, doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.top_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.header_distance = Inches(HEADER_DISTANCE)
    section.footer_distance = Inches(FOOTER_DISTANCE)
    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run(f"ROBOTICS PM OPERATING SYSTEM  |  {doc_type.upper()}"), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), size=24, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_run(p.add_run(subtitle), size=12, color=MUTED)
    values = [("Program", "[ENTER]"), ("Owner", "[ENTER]"), ("Version / date", "[ENTER]"), ("Approval state", "Draft — not approved")]
    for key, value in values:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(f"{key}: "), size=10, bold=True, color=NAVY)
        set_run(p.add_run(value), size=10)
    p = doc.add_paragraph(style="Evidence Note")
    set_run(p.add_run("Evidence-confidence key: [PB-H] public benchmark/high; [RBE-M] range estimate/medium; [SA-L] scenario assumption/low; [DC-L] derived calculation/low; [UPV] unverified production value."), size=9, color=MUTED)
    doc.core_properties.title = title
    doc.core_properties.subject = "Reusable robotics program management template"
    doc.core_properties.author = "Anthony Durham"
    doc.core_properties.comments = "Template fields and example values require human validation and approval."
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_text(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))


def add_table(doc: Document, number: int, title: str, headers: list[str], rows: list[list[str]], widths: list[int], evidence: str) -> None:
    caption = doc.add_paragraph(style="Table Title")
    set_run(caption.add_run(f"Table {number}. {title} — Evidence: {evidence}"), size=9, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        shade(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=8.5, bold=True, color=NAVY)
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            for run in cells[index].paragraphs[0].runs:
                set_run(run, size=8.5)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths)


def add_signature_table(doc: Document, number: int, title: str, roles: list[str], *, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    add_table(
        doc,
        number,
        title,
        ["Role", "Name", "Decision / signature", "Date"],
        [[role, "[ENTER]", "[APPROVE / HOLD + SIGN]", "[ENTER]"] for role in roles],
        [1800, 2160, 3600, 1800],
        "signed approval [UPV]; Confidence: unverified until signed",
    )


def program_brief() -> Document:
    doc = new_document("Robotics Program Brief", "A one-document baseline for outcomes, constraints, governance, and evidence", "program brief")
    add_heading(doc, "1. Decision frame")
    add_text(doc, "Decision requested: [ENTER the specific authorization, threshold, amount, and decision date].")
    add_text(doc, "Recommended disposition: [GO / CONDITIONAL GO / HOLD]. This recommendation is not approved until the signature table is complete.")
    add_table(doc, 1, "Outcome and boundary baseline", ["Dimension", "Approved statement", "Evidence / confidence"], [
        ["Customer outcome", "[ENTER measurable customer or operator outcome]", "[UPV] Unverified"],
        ["In scope", "[ENTER sites, robots, workflows, integrations, and support period]", "[SA-L] Low"],
        ["Out of scope", "[ENTER prohibited tasks, environments, actions, and claims]", "[SA-L] Low"],
        ["Success gate", "[ENTER metric, threshold, sample, window, and approver]", "[UPV] Unverified"],
    ], [1800, 5160, 2400], "template assumptions [SA-L] and approvals [UPV]; Confidence: low/unverified")
    add_heading(doc, "2. Integrated baseline")
    add_table(doc, 2, "Scope, schedule, cost, and benefit controls", ["Control", "Baseline", "Owner", "Evidence"], [
        ["Milestone / gate", "[ENTER date and predecessor]", "[ENTER]", "Approved schedule [UPV]"],
        ["Budget / TCO", "[ENTER currency, period, contingency, exclusions]", "[ENTER]", "Quote and model [UPV]"],
        ["Benefit", "[ENTER baseline, target, method, window]", "[ENTER]", "Measurement plan [UPV]"],
        ["Safety / security", "[ENTER acceptance rule and stop authority]", "[ENTER]", "Signed plan [UPV]"],
    ], [1800, 3600, 1560, 2400], "approved baselines required [UPV]; Confidence: unverified")
    add_heading(doc, "3. Governance and escalation")
    add_table(doc, 3, "Decision rights and cadence", ["Forum", "Cadence", "Owns", "Required input / output"], [
        ["Executive steering", "[ENTER]", "Investment, scope, risk acceptance", "Decision memo and signed log"],
        ["Program control", "Weekly", "Integrated plan, RAID, budget, dependencies", "Reconciled source registers"],
        ["Safety / security review", "Per gate", "Acceptance thresholds and exceptions", "Signed evidence package"],
        ["Operations review", "Weekly after launch", "Reliability, service, benefit, changes", "Telemetry and case evidence"],
    ], [1800, 1440, 2880, 3240], "recommended governance pattern [SA-L]; Confidence: low")
    add_heading(doc, "4. Top decisions and unknowns")
    add_table(doc, 4, "Decision and evidence queue", ["ID", "Decision / unknown", "Owner", "Due", "Evidence state"], [
        ["DEC-001", "[ENTER]", "[ENTER]", "[ENTER]", "[UPV] Pending"],
        ["DEC-002", "[ENTER]", "[ENTER]", "[ENTER]", "[UPV] Pending"],
        ["VAL-001", "[ENTER validation needed]", "[ENTER]", "[ENTER]", "[UPV] Pending"],
    ], [1200, 3960, 1560, 1320, 1320], "open decisions [UPV]; Confidence: unverified")
    add_signature_table(doc, 5, "Program baseline approval", ["Executive sponsor", "Program owner", "Safety owner", "Customer / operations owner"])
    return doc


def gate_memo() -> Document:
    doc = new_document("Stage-Gate Decision Memo", "A signed, evidence-based release decision with explicit conditions", "decision memo")
    add_heading(doc, "1. Recommendation")
    add_text(doc, "Gate: [G0 / G1 / G2 / G3 / G4]  |  Decision date: [ENTER]  |  Recommendation: [GO / CONDITIONAL GO / HOLD].")
    add_text(doc, "Decision rationale: [ENTER one paragraph that connects evidence to the recommendation and states residual exposure].")
    add_table(doc, 1, "Gate criteria assessment", ["Criterion", "Threshold", "Result", "State", "Evidence / confidence"], [
        ["Safety", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach signed result"],
        ["Technical performance", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach test report"],
        ["Site and operations", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach checklist"],
        ["Security and privacy", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach approval"],
        ["Financial", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach reconciled model"],
        ["Service readiness", "[ENTER]", "[ENTER]", "[PASS/HOLD]", "[UPV] Attach operating evidence"],
    ], [1800, 1800, 1800, 1200, 2760], "gate evidence [UPV]; Confidence: unverified until approvers sign")
    add_heading(doc, "2. Exceptions and conditions")
    add_table(doc, 2, "Conditions of approval", ["Condition", "Owner", "Due", "Closure evidence", "Stop trigger"], [
        ["[ENTER or state NONE]", "[ENTER]", "[ENTER]", "[ATTACH]", "[ENTER]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH]", "[ENTER]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH]", "[ENTER]"],
    ], [2520, 1560, 1320, 2160, 1800], "approval conditions [UPV]; Confidence: unverified")
    add_heading(doc, "3. Residual risk and reversibility")
    add_table(doc, 3, "Residual risk acceptance", ["Risk ID", "Exposure after mitigation", "Accountable acceptor", "Rollback / safe-state plan"], [
        ["[ENTER]", "[ENTER probability, impact, and consequence]", "[ENTER]", "[ENTER tested action and owner]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]"],
    ], [1200, 3120, 2160, 2880], "risk assessment and test evidence [UPV]; Confidence: unverified")
    add_heading(doc, "4. Decision")
    add_signature_table(doc, 4, "Gate decision record", ["Executive sponsor", "Program owner", "Safety owner", "Security / privacy owner", "Customer / operations owner"])
    return doc


def weekly_status() -> Document:
    doc = new_document("Weekly Executive Status", "A concise, evidence-aware program control and decision report", "weekly status")
    add_heading(doc, "1. Executive readout")
    add_table(doc, 1, "Executive health summary", ["Dimension", "RAG", "This-week evidence", "Next control"], [
        ["Outcome / benefit", "[R/A/G]", "[ENTER measured result + label]", "[ENTER]"],
        ["Schedule", "[R/A/G]", "[ENTER milestone variance]", "[ENTER]"],
        ["Budget / TCO", "[R/A/G]", "[ENTER actual/forecast/variance]", "[ENTER]"],
        ["Safety / quality", "[R/A/G]", "[ENTER test/incident evidence]", "[ENTER]"],
        ["Service / adoption", "[R/A/G]", "[ENTER telemetry/case/user evidence]", "[ENTER]"],
    ], [1800, 840, 4080, 2640], "weekly source registers [UPV]; Confidence: unverified until reconciled")
    add_heading(doc, "2. Accomplishments and next work")
    add_table(doc, 2, "Delivery evidence", ["Workstream", "Completed this week", "Next week", "Owner", "Evidence"], [
        ["Program", "[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH + label]"],
        ["Engineering / integration", "[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH + label]"],
        ["Site / operations", "[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH + label]"],
        ["Safety / security", "[ENTER]", "[ENTER]", "[ENTER]", "[ATTACH + label]"],
    ], [1560, 2520, 2520, 1320, 1440], "weekly records [UPV]; Confidence: unverified")
    add_heading(doc, "3. Exceptions and decisions")
    add_table(doc, 3, "Top RAID items", ["ID", "Exposure / trigger", "Response", "Owner / due", "State"], [
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]", "[OPEN/CLOSED]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]", "[OPEN/CLOSED]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]", "[OPEN/CLOSED]"],
    ], [1080, 2760, 2520, 1800, 1200], "RAID register [UPV]; Confidence: unverified")
    add_table(doc, 4, "Executive decisions required", ["Decision", "Options / recommendation", "Decision owner", "Needed by", "Impact if delayed"], [
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]", "[ENTER]"],
    ], [2040, 2760, 1680, 1320, 1560], "open decision log [UPV]; Confidence: unverified")
    add_heading(doc, "4. Reconciliation")
    add_table(doc, 5, "Source register checkpoint", ["Source", "Version / as-of", "Owner", "Reconciled?", "Exception"], [
        ["Schedule", "[ENTER]", "[ENTER]", "[YES/NO]", "[ENTER]"],
        ["Budget / TCO", "[ENTER]", "[ENTER]", "[YES/NO]", "[ENTER]"],
        ["RAID / decisions", "[ENTER]", "[ENTER]", "[YES/NO]", "[ENTER]"],
        ["Requirements / evidence", "[ENTER]", "[ENTER]", "[YES/NO]", "[ENTER]"],
    ], [2040, 1800, 1560, 1440, 2520], "named source records [UPV]; Confidence: unverified")
    return doc


def sat_plan() -> Document:
    doc = new_document("Site Acceptance Test Plan", "A controlled requirements-to-test protocol for robotics deployment", "site acceptance")
    add_heading(doc, "1. Test authority and boundary")
    add_text(doc, "Test objective: demonstrate only the approved requirements within the controlled site, task, payload, environmental, supervision, and recovery envelope.")
    add_text(doc, "Stop authority: any tester may call STOP. Safety Lead owns restart authorization after evidence review. No incomplete or missing test may be recorded as passed.")
    add_table(doc, 1, "Entry criteria", ["Entry criterion", "Owner", "Required evidence", "State"], [
        ["Approved hazard analysis and test envelope", "Safety Lead", "Signed document", "[PASS/HOLD]"],
        ["Correct hardware, software, model, and configuration identifiers", "Configuration Lead", "Serialized configuration record", "[PASS/HOLD]"],
        ["Site, network, privacy, and emergency readiness", "Site / Security Leads", "Signed readiness checklist", "[PASS/HOLD]"],
        ["Trained operators, observers, and stop authority", "Operations Lead", "Training record and briefing", "[PASS/HOLD]"],
    ], [3240, 1800, 3000, 1320], "entry evidence [UPV]; Confidence: unverified until signed")
    add_heading(doc, "2. Test matrix")
    add_table(doc, 2, "Acceptance test matrix", ["Test ID / requirement", "Method and sample", "Threshold", "Result", "Evidence / state"], [
        ["SAT-SAF-01 / REQ-SAF-001", "Induce approved communications loss; n=[ENTER]", "[ENTER stop envelope]", "[ENTER]", "[ATTACH] / [PASS/HOLD]"],
        ["SAT-OPS-01 / REQ-OPS-001", "Run eligible workflow; n=[ENTER]", "[ENTER success threshold]", "[ENTER]", "[ATTACH] / [PASS/HOLD]"],
        ["SAT-REC-01 / REQ-REC-001", "Inject recoverable fault; n=[ENTER]", "[ENTER recovery and escalation threshold]", "[ENTER]", "[ATTACH] / [PASS/HOLD]"],
        ["SAT-CYB-01 / REQ-CYB-001", "Inspect access, logs, revocation, and fail-safe", "All approved controls present", "[ENTER]", "[ATTACH] / [PASS/HOLD]"],
    ], [2040, 2520, 2040, 1200, 1560], "test requirements and results [UPV]; Confidence: unverified")
    add_heading(doc, "3. Execution record")
    add_table(doc, 3, "Run log and anomalies", ["Run", "Timestamp / configuration", "Observation", "Disposition", "Evidence ID"], [
        ["[ENTER]", "[ENTER]", "[ENTER]", "[PASS / FAIL / ABORT]", "[ENTER]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[PASS / FAIL / ABORT]", "[ENTER]"],
        ["[ENTER]", "[ENTER]", "[ENTER]", "[PASS / FAIL / ABORT]", "[ENTER]"],
    ], [1080, 2400, 2760, 1800, 1320], "observed test records [UPV]; Confidence: unverified")
    add_heading(doc, "4. Exit and handoff")
    add_table(doc, 4, "Exit criteria and residual actions", ["Criterion", "State", "Owner", "Due", "Closure evidence"], [
        ["All critical requirements passed", "[PASS/HOLD]", "Test Lead", "[ENTER]", "Signed RTM"],
        ["No open unacceptable safety/security exception", "[PASS/HOLD]", "Safety / Security", "[ENTER]", "Signed exception log"],
        ["Operations, support, spares, and rollback ready", "[PASS/HOLD]", "Operations Lead", "[ENTER]", "Handoff record"],
        ["Customer accepts the approved result and limitations", "[PASS/HOLD]", "Customer Owner", "[ENTER]", "Signed acceptance"],
    ], [3240, 1320, 1800, 1200, 1800], "exit evidence [UPV]; Confidence: unverified")
    add_signature_table(doc, 5, "Site acceptance decision", ["Test lead", "Safety owner", "Operations owner", "Customer owner"])
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    documents = {
        "ROBOTICS_PROGRAM_BRIEF.docx": program_brief(),
        "STAGE_GATE_DECISION_MEMO.docx": gate_memo(),
        "WEEKLY_EXECUTIVE_STATUS.docx": weekly_status(),
        "SITE_ACCEPTANCE_TEST_PLAN.docx": sat_plan(),
    }
    for name, doc in documents.items():
        path = OUT / name
        doc.save(path)
        print(f"Built {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
